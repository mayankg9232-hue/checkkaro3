import io
import os
import sys
from typing import Union, BinaryIO, Tuple, Dict, Any, List
from pypdf import PdfReader
import docx

MIN_PAGE_TEXT_CHARS = 20

def extract_text_from_pdf(
    file_input: Union[str, bytes, BinaryIO],
    return_metadata: bool = False
) -> Union[str, Tuple[str, Dict[str, Any]]]:
    """
    Extracts plain text from all pages of a PDF file (path, bytes, or file-like stream).
    
    Robust Scanning Rules:
    - Iterates through EVERY page without early termination or page limits.
    - If a page contains an image or < 20 readable characters, skips visual content silently,
      notes the page in low_text_pages, and continues scanning all subsequent pages.
    - Concatenates all readable page texts in original page order.
    - Raises ValueError only if the entire document is empty or completely unreadable.
    """
    try:
        if isinstance(file_input, str):
            if not os.path.exists(file_input):
                raise FileNotFoundError(f"File not found: '{file_input}'")
            with open(file_input, "rb") as f:
                stream = io.BytesIO(f.read())
        elif isinstance(file_input, bytes):
            stream = io.BytesIO(file_input)
        else:
            # File-like object (e.g. Streamlit UploadedFile)
            if hasattr(file_input, "getvalue"):
                stream = io.BytesIO(file_input.getvalue())
            elif hasattr(file_input, "read"):
                stream = io.BytesIO(file_input.read())
                if hasattr(file_input, "seek"):
                    file_input.seek(0)
            else:
                stream = file_input

        reader = PdfReader(stream)
        total_pages = len(reader.pages)
        if total_pages == 0:
            raise ValueError("The PDF document contains 0 pages or is empty.")

        extracted_pages: List[str] = []
        low_text_pages: List[int] = []

        for i, page in enumerate(reader.pages):
            page_num = i + 1
            try:
                raw_text = page.extract_text() or ""
            except Exception:
                # If a specific page contains corrupted image streams, skip gracefully
                raw_text = ""

            clean_text = raw_text.strip()
            if len(clean_text) < MIN_PAGE_TEXT_CHARS:
                # Image-only, blank, or trivial page: note and skip without error
                low_text_pages.append(page_num)
            else:
                extracted_pages.append(clean_text)

        full_text = "\n\n".join(extracted_pages).strip()

        if not full_text:
            raise ValueError("The PDF document does not contain readable text (it may be scanned/image-only or blank).")

        metadata = {
            "total_pages": total_pages,
            "readable_pages": len(extracted_pages),
            "low_text_pages": low_text_pages,
            "format": "PDF"
        }

        if return_metadata:
            return full_text, metadata
        return full_text

    except ValueError as ve:
        raise ve
    except Exception as e:
        raise RuntimeError(f"Corrupted or invalid PDF file: {str(e)}")

def extract_text_from_docx(
    file_input: Union[str, bytes, BinaryIO],
    return_metadata: bool = False
) -> Union[str, Tuple[str, Dict[str, Any]]]:
    """
    Extracts plain text from all paragraphs and tables of a DOCX file.
    Silently skips non-text shapes/drawings and concatenates all textual content.
    """
    try:
        if isinstance(file_input, str):
            if not os.path.exists(file_input):
                raise FileNotFoundError(f"File not found: '{file_input}'")
            with open(file_input, "rb") as f:
                stream = io.BytesIO(f.read())
        elif isinstance(file_input, bytes):
            stream = io.BytesIO(file_input)
        else:
            if hasattr(file_input, "getvalue"):
                stream = io.BytesIO(file_input.getvalue())
            elif hasattr(file_input, "read"):
                stream = io.BytesIO(file_input.read())
                if hasattr(file_input, "seek"):
                    file_input.seek(0)
            else:
                stream = file_input

        doc = docx.Document(stream)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs).strip()
        if not full_text:
            raise ValueError("The DOCX document is empty.")

        metadata = {
            "total_pages": 1,
            "readable_pages": 1,
            "low_text_pages": [],
            "format": "DOCX"
        }

        if return_metadata:
            return full_text, metadata
        return full_text

    except ValueError as ve:
        raise ve
    except Exception as e:
        raise RuntimeError(f"Corrupted or invalid DOCX file: {str(e)}")

def get_text(
    uploaded_file: Union[str, bytes, BinaryIO],
    filename: str = None,
    return_metadata: bool = False
) -> Union[str, Tuple[str, Dict[str, Any]]]:
    """
    Primary document extraction interface. Scans entire PDF/DOCX documents,
    gracefully skipping unreadable/image pages while preserving all readable content.
    """
    if uploaded_file is None:
        raise ValueError("No file was provided for text extraction.")

    name = filename or getattr(uploaded_file, "name", "")
    if not name and isinstance(uploaded_file, str):
        name = uploaded_file

    lower_name = name.lower()
    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file, return_metadata=return_metadata)
    elif lower_name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file, return_metadata=return_metadata)
    else:
        ext = os.path.splitext(name)[1] if name else "unknown"
        raise ValueError(f"Unsupported file format '{ext}'. Only PDF (.pdf) and DOCX (.docx) files are supported.")

def extract_text_with_metadata(
    uploaded_file: Union[str, bytes, BinaryIO],
    filename: str = None
) -> Tuple[str, Dict[str, Any]]:
    """
    Convenience function that always returns (extracted_text, metadata_dict).
    """
    return get_text(uploaded_file, filename=filename, return_metadata=True)

# Alias for backward compatibility across modules
extract_text = get_text

if __name__ == "__main__":
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")

    if len(sys.argv) > 1:
        target_path = sys.argv[1]
        try:
            print(f"Extracting text from: {target_path}")
            result_text, meta = extract_text_with_metadata(target_path)
            print("=" * 60)
            print(result_text)
            print("=" * 60)
            print(f"Extraction successful! Characters: {len(result_text)}, Metadata: {meta}")
        except Exception as err:
            print(f"Extraction Error: {err}", file=sys.stderr)
            sys.exit(1)
    else:
        print("Usage: python logic/extract_text.py <path_to_pdf_or_docx>")

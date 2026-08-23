import io
import os
import docx
from pypdf import PdfWriter
from pypdf.generic import NameObject, create_string_object, DictionaryObject, DecodedStreamObject

os.makedirs("samples", exist_ok=True)

# 1. Create sample Income Tax Notice docx
doc = docx.Document()
doc.add_heading("INCOME TAX DEPARTMENT - GOVERNMENT OF INDIA", level=1)
doc.add_paragraph("Intimation under Section 143(1) of the Income-tax Act, 1961")
doc.add_paragraph("Assessment Year: 2025-26 | PAN: ABCDE1234F | Notice Date: 15-Aug-2026")
doc.add_paragraph("Dear Taxpayer, Your return of income for AY 2025-26 has been processed at the Centralized Processing Centre (CPC), Bengaluru.")
doc.add_paragraph("Upon processing, there is an outstanding net demand payable as summarized below:")

table = doc.add_table(rows=3, cols=2)
table.rows[0].cells[0].text = "Particulars"
table.rows[0].cells[1].text = "Amount (INR)"
table.rows[1].cells[0].text = "Tax Payable under Section 140A"
table.rows[1].cells[1].text = "₹ 4,000"
table.rows[2].cells[0].text = "Interest under Section 234B/234C"
table.rows[2].cells[1].text = "₹ 500"

doc.add_paragraph("\nACTION REQUIRED:")
doc.add_paragraph("1. Log in to the e-Filing portal (www.incometax.gov.in) with your PAN.")
doc.add_paragraph("2. Navigate to 'e-File' > 'Response to Outstanding Demand'.")
doc.add_paragraph("3. Either 'Agree with Demand' and make payment via Challan ITNS 280 / e-Pay Tax within 30 days of receipt, or 'Disagree with Demand' and upload supporting tax deduction proof.")
doc.add_paragraph("REQUIRED DOCUMENTS:")
doc.add_paragraph("- Form 16 / 16A TDS certificates\n- Bank statement showing advance tax deduction\n- Original ITR-V filing acknowledgement")
doc.add_paragraph("IMPORTANT WARNING / RISKS:")
doc.add_paragraph("- Failure to pay the demand within 30 days will attract additional interest under Section 220(2) at 1% per month.")
doc.add_paragraph("- Recovery proceedings may be initiated by the jurisdictional Assessing Officer (AO).")

doc.save("samples/sample_it_notice.docx")
print("Created samples/sample_it_notice.docx")

# 2. Create sample PDF
writer = PdfWriter()
page = writer.add_blank_page(width=595, height=842) # A4 size

stream = DecodedStreamObject()
pdf_text_commands = """
BT
/F1 16 Tf
50 780 Td
(GOVERNMENT OF KARNATAKA - REVENUE DEPARTMENT) Tj
0 -24 Td
/F1 13 Tf
(Notice of Property Khata Verification and Transfer) Tj
0 -20 Td
/F1 10 Tf
(Application Ref: KHA-2026-98124 | Date: 18-Aug-2026) Tj
0 -25 Td
(Subject: Mandatory biometric and original deed verification for e-Khata issuance.) Tj
0 -20 Td
(Applicant: Smt. Ananya Rao | Ward: 174 HSR Layout, Bengaluru) Tj
0 -30 Td
(INSTRUCTIONS & ACTION REQUIRED:) Tj
0 -15 Td
(1. Visit the local BBMP Assistant Revenue Officer [ARO] office within 21 days.) Tj
0 -15 Td
(2. Submit physical copies of Registered Sale Deed, latest property tax receipt, and EC.) Tj
0 -15 Td
(3. Provide biometric verification at the citizen service counter.) Tj
0 -25 Td
(REQUIRED DOCUMENTS:) Tj
0 -15 Td
(- Registered Title Deed / Sale Deed) Tj
0 -15 Td
(- Latest Property Tax Paid Receipt for current financial year) Tj
0 -15 Td
(- Encumbrance Certificate [EC] for past 15 years from Sub-Registrar) Tj
0 -15 Td
(- Aadhaar Card and PAN Card of the applicant) Tj
0 -25 Td
(WARNING & PENALTY:) Tj
0 -15 Td
(Failure to complete verification within 21 days will lead to rejection of e-Khata application.) Tj
0 -15 Td
(Unverified properties cannot be transferred, registered, or mortgaged.) Tj
ET
"""
stream.set_data(pdf_text_commands.encode("utf-8"))
page[NameObject("/Contents")] = stream

fonts = DictionaryObject()
font1 = DictionaryObject()
font1[NameObject("/Type")] = NameObject("/Font")
font1[NameObject("/Subtype")] = NameObject("/Type1")
font1[NameObject("/BaseFont")] = NameObject("/Helvetica")
fonts[NameObject("/F1")] = font1

res = DictionaryObject()
res[NameObject("/Font")] = fonts
page[NameObject("/Resources")] = res

with open("samples/sample_khata_notice.pdf", "wb") as f:
    writer.write(f)

print("Created samples/sample_khata_notice.pdf")

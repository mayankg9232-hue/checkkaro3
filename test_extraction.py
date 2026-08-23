import io
import docx
from logic.extract_text import extract_text, extract_text_from_pdf, extract_text_from_docx

def create_sample_docx():
    doc = docx.Document()
    doc.add_heading('Income Tax Intimation Section 143(1)', level=1)
    doc.add_paragraph('Dear Taxpayer, Please note that your return for Assessment Year 2025-26 has been processed.')
    doc.add_paragraph('Outstanding demand amount is Rs. 4,500 payable within 30 days of receipt of this notice.')
    
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = 'Item'
    table.rows[0].cells[1].text = 'Value'
    table.rows[1].cells[0].text = 'Total Tax Payable'
    table.rows[1].cells[1].text = 'INR 4500'
    
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes

def test():
    print('Testing DOCX extraction...')
    docx_bytes = create_sample_docx()
    text_docx = extract_text(docx_bytes, filename='notice.docx')
    print('Extracted text sample:\n' + '-'*40)
    print(text_docx)
    print('-'*40)
    assert 'Income Tax Intimation' in text_docx
    assert '4,500' in text_docx or '4500' in text_docx
    print('DOCX extraction test: PASSED\n')
    
    print('Testing unsupported file format...')
    try:
        extract_text(b'sample', filename='test.txt')
        print('Failed: should have raised ValueError')
    except ValueError as e:
        print('Unsupported format handled correctly:', e)
        
    print('\nAll extraction unit tests PASSED successfully!')

if __name__ == '__main__':
    test()
